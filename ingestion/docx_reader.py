"""
DOCX Reader Module
Handles extraction of text from Microsoft Word documents.
"""

from pathlib import Path
from typing import Optional
from docx import Document
from loguru import logger


class DOCXReader:
    """Extracts text content from DOCX files."""
    
    def __init__(self):
        """Initialize DOCX reader."""
        pass
    
    def extract_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path.name}")
            
            doc = Document(str(file_path))
            text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            logger.info(f"Successfully extracted {len(full_text)} characters from {file_path.name}")
            
            return full_text if full_text.strip() else None
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path.name}: {e}")
            return None
    
    def extract_metadata(self, file_path: Path) -> dict:
        """
        Extract metadata from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing metadata
        """
        try:
            doc = Document(str(file_path))
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created,
                'modified': core_props.modified
            }
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path.name}: {e}")
            return {}
